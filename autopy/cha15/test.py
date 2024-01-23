import docx

# pdfFileObj = open('meetingminutes.pdf','rb')
# pdfReader = PyPDF2.PdfReader(pdfFileObj)
# print(len(pdfReader.pages))
# pageObj = pdfReader.pages[0]
# print(pageObj.extract_text())
# pdfFileObj.close()

# with open('encrypted.pdf','rb') as fp:
#     pdf_reader = PyPDF2.PdfReader(fp)
#     print(pdf_reader.is_encrypted)
#     try:
#         page_obj = pdf_reader.pages[0]
#         print(page_obj.extract_text())
#     except:
#         print("can't read!")
#     # pdf_reader = PyPDF2.PdfReader(fp)
#     pdf_reader.decrypt('rosebud')
#     page_obj = pdf_reader.pages[0]
#     print(page_obj.extract_text())

# pdf1File = open('meetingminutes.pdf','rb')
# pdf2File = open('meetingminutes2.pdf','rb')

# pdf1Reader = PyPDF2.PdfReader(pdf1File)
# pdf2Reader = PyPDF2.PdfReader(pdf2File)

# pdf_writer = PyPDF2.PdfWriter()

# for page in pdf1Reader.pages:
#     pdf_writer.add_page(page)
# for page in pdf2Reader.pages:
#     pdf_writer.add_page(page)

# pdfOutputFile = open('combinedminutes.pdf','wb')
# pdf_writer.write(pdfOutputFile)
# pdfOutputFile.close()
# pdf1File.close()
# pdf2File.close()
# with open('meetingminutes.pdf','rb') as fp :
#     pdf_reader = PyPDF2.PdfReader(fp)
#     page_obj = pdf_reader.pages[0]

#     page_obj.rotate(90)

#     pdf_writer = PyPDF2.PdfWriter()
#     fw = open('rotared.pdf','wb')
#     pdf_writer.add_page(page_obj)
#     for page in pdf_reader.pages[1:]:
#         pdf_writer.add_page(page.rotate(90))
#     pdf_writer.write(fw)

#     fw.close()

# with open('meetingminutes.pdf','rb') as minutesFile :
#     pdfReader = PyPDF2.PdfReader(minutesFile)
#     minutesFirstPage = pdfReader.pages[0]

#     with  open('watermark.pdf','rb') as waterMarkFile :
#         pdfWatermarkReader = PyPDF2.PdfReader(waterMarkFile)
#         # minutesFirstPage.merge_page(pdfWatermarkReader.pages[0])
#         watermarkFirstPage = pdfWatermarkReader.pages[0]
#         watermarkFirstPage.merge_page(minutesFirstPage)

#         pdfWriter = PyPDF2.PdfWriter()
#         # pdfWriter.add_page(minutesFirstPage)
#         pdfWriter.add_page(watermarkFirstPage)

#     for page in pdfReader.pages[1:]:
#         pdfWriter.add_page(page)

#     with open('merged.pdf','wb') as result:
#         pdfWriter.write(result)


# pdfFile = open('meetingminutes.pdf','rb')
# pdfReader = PyPDF2.PdfReader(pdfFile)

# pdfWriter = PyPDF2.PdfWriter()

# for page in pdfReader.pages:
#     pdfWriter.add_page(page)

# pdfWriter.encrypt('swordfish')

# resultPdf = open('encryptedminutes.pdf','wb')
# pdfWriter.write(resultPdf)

# resultPdf.close()
# pdfFile.close()

# doc = docx.Document('demo.docx')
# print(len(doc.paragraphs))

# import readDoc

# print('  '+readDoc.getText('demo.docx'))

# import docx
# doc = docx.Document('demo.docx')
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[0].style)

# doc.paragraphs[0].style = 'Normal'
# print(doc.paragraphs[1].text)

# import docx
# doc = docx.Document('test1.docx')
# # print(doc.paragraphs[0].text)
# doc.paragraphs[0].style = 'Title'
# doc.paragraphs[1].runs[0].underline = True
# # doc.paragraphs[2].runs[0].style = 'TitleChar'
# doc.paragraphs[2].runs[0].font.strike = True

# doc.save('new.docx')


# import docx
# doc = docx.Document()
# doc.add_paragraph('Hello, world!','Title')
# para1 = doc.add_paragraph('This is the second paragraph')
# para2 = doc.add_paragraph('This is anothor paragraph')

# para1.add_run('This text is added to the 2nd paragraph').font.strike = True

# doc.save('multiparagraphs.docx')

# import docx
# doc = docx.Document()
# # for num in range(5):
# #     doc.add_heading('Heading'+str(num),num)

# # doc.save('headinga.docx')
# doc.add_paragraph('This is on the 1st page.')
# doc.paragraphs[0].runs[0].add_break()# docx.enum.text.WD_BREAK.PAGE)
# doc.add_paragraph('This is on the 2nd page.')
# doc.add_picture('zophie.png') #,width=docx.shared.Inches(1),height=docx.shared.Cm(4))
# doc.save('twoLineOneParaWithPic.docx')

import docx,win32com.client
wordFileName = 'neword.docx'
pdfFileName = 'transToPdf.pdf'

doc = docx.Document()
# TODO: edit docx file
para1 = doc.add_paragraph('1st Paragraph','Title')
para1.runs[0].add_break()
para1.add_run('index').blod = True
doc.add_paragraph('2nd para').runs[0].font.size = docx.shared.Pt(20)

doc.save(wordFileName)

wdFormatPDF = 17
wordObj = win32com.client.Dispatch('Word.Application')
docObj = wordObj.Documents.Open(wordFileName)
docObj.SaveAs(pdfFileName,FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()

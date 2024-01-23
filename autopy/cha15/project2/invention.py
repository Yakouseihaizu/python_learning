import docx

doc = docx.Document('invention.docx')
guests = []
with open('guest.txt') as fp:
    guestl = fp.readlines()
    for guest in guestl:
        if guest.endswith('\n'):
            guest = guest[:-1]
        guests.append(guest)

for guest in guests:
    paragraph1 = doc.add_paragraph('It would be a pleasure to have the company of')
    paragraph1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    paragraph1.style = 'sample1'

    paragraph2 = doc.add_paragraph(guest)
    paragraph2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    paragraph2.style = 'sample2'
    
    paragraph3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    paragraph3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    paragraph3.style = 'sample1'
    
    paragraph4 = doc.add_paragraph('April 1st')
    #  paragraph1.style = ''
    paragraph4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
   
    paraObj = doc.add_paragraph("at 7 o'clock")
    paraObj.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    paraObj.style = 'sample1'
    paraObj.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    
doc.save('invention.docx')